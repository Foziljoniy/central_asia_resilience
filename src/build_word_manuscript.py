"""Build a clean Word manuscript from the approved Markdown manuscript."""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "full_manuscript_v2_clean.md"
OUTPUT = ROOT / "manuscript" / "final" / "Foziljon_Alisherov_Remittances_Shocks.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.strip())
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, width_inches=7.0):
    widths = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    cell_width = Inches(width_inches / widths)
    font_size = 7.5 if widths >= 10 else 8.5
    for i, row in enumerate(rows):
        for j in range(widths):
            value = row[j] if j < len(row) else ""
            table.cell(i, j).width = cell_width
            set_cell_text(table.cell(i, j), value, bold=(i == 0), font_size=font_size)
            if i == 0:
                shade(table.cell(i, j), "D9EAF0")
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_tr_pr.append(repeat)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def readable_math(text):
    replacements = {
        r"\beta_0": "β₀", r"\beta_1": "β₁", r"\beta_2": "β₂", r"\beta_3": "β₃",
        r"\gamma": "γ", r"\delta_c": "δ_c", r"\tau_t": "τ_t",
        r"\varepsilon_{ict}": "ε_ict", r"\times": "×",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.replace(r"\(", "").replace(r"\)", "").replace("_{ict}", "_ict")


def add_markdown_paragraph(doc, text):
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = readable_math(text.replace("**", ""))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    for part in re.split(r"(`[^`]+`|\*[^*]+\*)", text):
        if not part:
            continue
        run = p.add_run(part.strip("`") if part.startswith("`") else part.strip("*"))
        run.font.name = "Aptos"
        run.font.size = Pt(10.5)
        if part.startswith("*") and not part.startswith("**"):
            run.italic = True
    return p


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line == r"\[":
            equation_parts = []
            i += 1
            while i < len(lines) and lines[i].strip() != r"\]":
                equation_parts.append(lines[i].strip())
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(readable_math(" ".join(equation_parts)))
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            run.italic = True
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.name = "Aptos Display"
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x0E, 0x4F, 0x5C)
            first_title = False
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.name = "Aptos Display"
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x0E, 0x4F, 0x5C)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.name = "Aptos Display"
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x2A, 0x6F, 0x7B)
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.add_run(line[2:].strip()).font.size = Pt(10.5)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw = [x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", x) for x in raw):
                    rows.append(raw)
                i += 1
            if rows:
                add_table(doc, rows)
            continue
        add_markdown_paragraph(doc, line)
        i += 1

    # Append the approved manuscript tables and final figures so the Word file
    # contains the same supporting material as the final LaTeX package.
    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width = Inches(11)
    landscape.page_height = Inches(8.5)
    landscape.top_margin = Inches(0.55)
    landscape.bottom_margin = Inches(0.55)
    landscape.left_margin = Inches(0.5)
    landscape.right_margin = Inches(0.5)
    add_heading = doc.add_paragraph()
    add_heading.style = doc.styles["Heading 1"]
    add_heading.add_run("Tables")
    table_dir = ROOT / "manuscript" / "tables_v2"
    for table_index, table_path in enumerate(sorted(table_dir.glob("table_*.md"))):
        if table_index:
            doc.add_page_break()
        caption = doc.add_paragraph()
        caption.paragraph_format.space_before = Pt(8)
        caption.paragraph_format.space_after = Pt(4)
        run = caption.add_run(table_path.stem.replace("_", " ").title())
        run.bold = True
        rows = []
        table_lines = table_path.read_text(encoding="utf-8").splitlines()
        for raw in table_lines:
            if raw.strip().startswith("|"):
                cells = [x.strip() for x in raw.strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", x) for x in cells):
                    rows.append(cells)
        if rows:
            add_table(doc, rows, width_inches=10.0)

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width = Inches(8.5)
    portrait.page_height = Inches(11)
    portrait.top_margin = Inches(0.8)
    portrait.bottom_margin = Inches(0.8)
    portrait.left_margin = Inches(0.9)
    portrait.right_margin = Inches(0.9)
    fig_heading = doc.add_paragraph()
    fig_heading.style = doc.styles["Heading 1"]
    fig_heading.add_run("Figures")
    figures = [
        ("figure_19_kyrgyzstan_adjusted_four_groups_v2.png", "Figure 19. Kyrgyzstan adjusted food-insecurity predictions by remittance and shock status."),
        ("figure_25_uzbekistan_broad_shock_predictions.png", "Figure 25. Uzbekistan adjusted food-insecurity predictions by remittance and verified-shock status."),
        ("figure_26_revised_standardized_interactions.png", "Figure 26. Standardized remittance-shock interaction associations."),
        ("figure_24_kazakhstan_benchmark_with_ci.png", "Figure 24. Kazakhstan annual food-insecurity benchmark."),
    ]
    for figure_index, (filename, caption_text) in enumerate(figures):
        image_path = ROOT / "outputs" / "figures" / filename
        if not image_path.exists():
            continue
        if figure_index:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(6.2))
        cp = doc.add_paragraph(caption_text)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)

    # Add explicit title-page metadata if the source front matter changes later.
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Foziljon Alisherov and Mukhayyo Djuraeva | Remittances and household shocks")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
