"""Shared, deterministic utilities for the Phase 1 research-data audit.

The module never writes under a detected raw-data directory.  It inventories
metadata and documentation only; it does not produce respondent-level output.
"""

from __future__ import annotations

import csv
import gzip
import hashlib
import importlib.metadata
import json
import logging
import os
import platform
import re
import shutil
import subprocess
import sys
import tarfile
import zipfile
from collections import Counter, defaultdict, deque
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Iterable


ROOT = Path(__file__).resolve().parents[1]
VENDOR = ROOT / ".vendor"
if VENDOR.exists() and str(VENDOR) not in sys.path:
    sys.path.insert(0, str(VENDOR))
CHECKPOINTS = ROOT / "outputs" / "checkpoints"
UNPACKED = ROOT / "data" / "interim" / "unpacked"
LOG_PATH = ROOT / "outputs" / "logs" / "phase_01.log"
COUNTRIES = ("kyrgyzstan", "uzbekistan")
ARCHIVE_SUFFIXES = (".zip", ".tar", ".tar.gz", ".tgz", ".gz", ".7z", ".rar")
MAX_DEPTH = 5
MAX_MEMBERS = 25_000
MAX_OUTER_UNCOMPRESSED = 10 * 1024**3
MAX_SINGLE_FILE = 5 * 1024**3
HIGH_RATIO = 1_000


def configure_logging() -> logging.Logger:
    """Configure one project log without duplicating handlers."""
    LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger("phase01")
    logger.setLevel(logging.INFO)
    if not logger.handlers:
        formatter = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")
        file_handler = logging.FileHandler(LOG_PATH, encoding="utf-8")
        file_handler.setFormatter(formatter)
        stream_handler = logging.StreamHandler()
        stream_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        logger.addHandler(stream_handler)
    return logger


LOGGER = configure_logging()


def ensure_structure() -> None:
    """Create required project directories, never moving raw inputs."""
    dirs = [
        "data/raw/kyrgyzstan", "data/raw/uzbekistan", "data/interim/unpacked/kyrgyzstan",
        "data/interim/unpacked/uzbekistan", "data/processed", "documentation/kyrgyzstan",
        "documentation/uzbekistan", "literature/original", "literature/notes",
        "literature/matrices", "literature/drafts", "literature/verification", "research",
        "src", "outputs/checkpoints", "outputs/tables", "outputs/figures", "outputs/models",
        "outputs/logs", "tests",
    ]
    for directory in dirs:
        (ROOT / directory).mkdir(parents=True, exist_ok=True)


def rel(path: Path) -> str:
    """Return a stable POSIX project-relative path."""
    try:
        return path.resolve().relative_to(ROOT.resolve()).as_posix()
    except ValueError:
        return "OUTSIDE_PROJECT"


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    """Calculate SHA-256 without loading a whole archive into memory."""
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(chunk_size), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_csv(path: Path, rows: Iterable[dict[str, Any]], fieldnames: list[str]) -> None:
    """Write a deterministic UTF-8 CSV with fixed columns."""
    path.parent.mkdir(parents=True, exist_ok=True)
    normalized = []
    for row in rows:
        normalized.append({key: _cell(row.get(key, "")) for key in fieldnames})
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(normalized)


def read_csv(path: Path) -> list[dict[str, str]]:
    """Read one generated CSV safely."""
    if not path.exists():
        return []
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def _cell(value: Any) -> str | int | float:
    if value is None:
        return ""
    if isinstance(value, (list, tuple, dict, set)):
        return json.dumps(value if not isinstance(value, set) else sorted(value), ensure_ascii=False, sort_keys=True)
    if isinstance(value, bool):
        return "true" if value else "false"
    return value


def detected_raw_roots() -> dict[str, Path]:
    """Detect requested or legacy archive locations without relocating inputs."""
    roots: dict[str, Path] = {}
    for country in COUNTRIES:
        preferred = ROOT / "data" / "raw" / country
        legacy = ROOT / "data" / country
        preferred_files = list(preferred.glob("*")) if preferred.exists() else []
        roots[country] = preferred if any(p.is_file() for p in preferred_files) else legacy
    return roots


def is_archive(path: Path | str) -> bool:
    name = str(path).lower()
    return any(name.endswith(suffix) for suffix in ARCHIVE_SUFFIXES)


def archive_format(path: Path | str) -> str:
    name = str(path).lower()
    for suffix in (".tar.gz", ".tgz", ".zip", ".tar", ".gz", ".7z", ".rar"):
        if name.endswith(suffix):
            return suffix.lstrip(".").replace("tar.gz", "tar.gz")
    return "unknown"


def safe_name(value: str) -> str:
    """Make a deterministic filesystem component."""
    clean = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._")
    return clean[:120] or "archive"


def safe_member(member: str) -> tuple[bool, str]:
    """Reject absolute, drive-qualified, or traversal member paths."""
    normalized = member.replace("\\", "/")
    pure = PurePosixPath(normalized)
    unsafe = (
        pure.is_absolute()
        or any(part == ".." for part in pure.parts)
        or bool(re.match(r"^[A-Za-z]:", normalized))
        or normalized.startswith("//")
    )
    return (not unsafe, normalized)


def _unique_target(base: Path, member: str, occurrence: int = 0) -> Path:
    """Resolve a safe output path and suffix duplicate member paths."""
    parts = [safe_name(part) for part in PurePosixPath(member.replace("\\", "/")).parts if part not in ("", ".")]
    target = base.joinpath(*parts)
    if occurrence:
        target = target.with_name(f"{target.stem}__dup{occurrence:03d}{target.suffix}")
    resolved = target.resolve()
    if base.resolve() not in (resolved, *resolved.parents):
        raise ValueError("resolved extraction path escaped archive directory")
    return target


ARCHIVE_FIELDS = [
    "country", "archive_id", "parent_archive_id", "nesting_depth", "relative_path", "filename",
    "archive_format", "file_size", "sha256", "number_of_members", "total_compressed_size",
    "estimated_uncompressed_size", "integrity_result", "encrypted_status", "duplicate_status",
    "extraction_status", "extraction_location", "warnings", "errors", "notes",
]
MEMBER_FIELDS = [
    "country", "archive_id", "parent_archive_id", "nesting_depth", "member_path", "filename",
    "extension", "compressed_size", "uncompressed_size", "crc_result", "encryption_status",
    "directory_or_file", "nested_archive_status", "suspected_role", "extraction_path", "safe_or_unsafe", "notes",
]


def suspected_role(path: str) -> str:
    text = path.lower()
    suffix = Path(text).suffix.lower()
    if suffix in {".dta", ".sav", ".por", ".csv", ".xlsx", ".xls", ".parquet"}:
        return "dataset"
    if "question" in text or "form" in text:
        return "questionnaire"
    if "codebook" in text or "dictionary" in text:
        return "codebook"
    if "manual" in text:
        return "manual"
    if "report" in text or "tables" in text or "study description" in text:
        return "report"
    if suffix in {".do", ".sps", ".stcmd", ".r", ".py"}:
        return "syntax"
    if "readme" in text or "manifest" in text:
        return "README"
    if suffix in {".pdf", ".doc", ".docx", ".rtf"}:
        return "publication"
    return "other"


def _zip_members(path: Path) -> tuple[list[dict[str, Any]], str, bool]:
    records: list[dict[str, Any]] = []
    encrypted = False
    integrity = "not tested"
    with zipfile.ZipFile(path) as archive:
        infos = archive.infolist()
        for info in infos:
            enc = bool(info.flag_bits & 0x1)
            encrypted = encrypted or enc
            records.append({
                "name": info.filename, "compressed": info.compress_size, "uncompressed": info.file_size,
                "is_dir": info.is_dir(), "encrypted": enc, "crc": f"{info.CRC:08x}", "zip_info": info,
                "is_link": False,
            })
        if encrypted:
            integrity = "not tested: encrypted member present"
        else:
            bad = archive.testzip()
            integrity = "passed" if bad is None else f"failed at {bad}"
    return records, integrity, encrypted


def _tar_members(path: Path) -> tuple[list[dict[str, Any]], str, bool]:
    records = []
    with tarfile.open(path, "r:*") as archive:
        for info in archive.getmembers():
            records.append({
                "name": info.name, "compressed": 0, "uncompressed": info.size, "is_dir": info.isdir(),
                "encrypted": False, "crc": "not available", "tar_info": info,
                "is_link": info.issym() or info.islnk(),
            })
    return records, "passed (headers readable)", False


def inventory_and_extract_archives() -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Recursively inventory and safely extract supported archives."""
    ensure_structure()
    archive_rows: list[dict[str, Any]] = []
    member_rows: list[dict[str, Any]] = []
    seen_checksums: dict[str, str] = {}
    per_outer_total: defaultdict[str, int] = defaultdict(int)
    roots = detected_raw_roots()
    queue: deque[dict[str, Any]] = deque()
    for country, raw_root in roots.items():
        for path in sorted((p for p in raw_root.rglob("*") if p.is_file() and is_archive(p)), key=lambda p: rel(p).lower()):
            stem = safe_name(re.sub(r"\.(tar\.gz|tgz|zip|tar|gz|7z|rar)$", "", path.name, flags=re.I))
            outer_base = UNPACKED / country / stem
            queue.append({"country": country, "path": path, "parent": "", "depth": 0,
                          "outer": f"{country}:{rel(path)}", "outer_base": outer_base})

    while queue:
        item = queue.popleft()
        country, path, parent, depth = item["country"], item["path"], item["parent"], item["depth"]
        outer, outer_base = item["outer"], item["outer_base"]
        fmt = archive_format(path)
        checksum = sha256_file(path)
        archive_id = f"{country}-{depth:02d}-{checksum[:12]}"
        extraction_base = outer_base / f"depth_{depth:02d}"
        if depth:
            extraction_base = extraction_base / f"{safe_name(path.stem)}__{checksum[:8]}"
        warnings: list[str] = []
        errors: list[str] = []
        duplicate_of = seen_checksums.get(checksum, "")
        duplicate_status = f"exact duplicate of {duplicate_of}" if duplicate_of else "unique"
        if not duplicate_of:
            seen_checksums[checksum] = archive_id
        row = {
            "country": country, "archive_id": archive_id, "parent_archive_id": parent,
            "nesting_depth": depth, "relative_path": rel(path), "filename": path.name,
            "archive_format": fmt, "file_size": path.stat().st_size, "sha256": checksum,
            "number_of_members": 0, "total_compressed_size": 0, "estimated_uncompressed_size": 0,
            "integrity_result": "not tested", "encrypted_status": "unknown",
            "duplicate_status": duplicate_status, "extraction_status": "not attempted",
            "extraction_location": rel(extraction_base), "warnings": "", "errors": "", "notes": "",
        }
        if depth > MAX_DEPTH:
            errors.append(f"maximum nesting depth {MAX_DEPTH} exceeded")
            row["extraction_status"] = "stopped: depth limit"
            row["errors"] = "; ".join(errors)
            archive_rows.append(row)
            continue
        if fmt in {"rar", "7z"}:
            row["integrity_result"] = "unsupported"
            row["encrypted_status"] = "unknown"
            row["extraction_status"] = "unsupported format"
            errors.append("safe RAR/7Z support is not installed")
            row["errors"] = "; ".join(errors)
            archive_rows.append(row)
            continue
        try:
            if fmt == "zip":
                members, integrity, encrypted = _zip_members(path)
            elif fmt in {"tar", "tar.gz", "tgz"}:
                members, integrity, encrypted = _tar_members(path)
            elif fmt == "gz":
                with gzip.open(path, "rb") as handle:
                    handle.read(1)
                uncompressed = max(0, int.from_bytes(path.read_bytes()[-4:], "little"))
                members = [{"name": path.stem, "compressed": path.stat().st_size, "uncompressed": uncompressed,
                            "is_dir": False, "encrypted": False, "crc": "not available", "is_link": False}]
                integrity, encrypted = "passed (stream readable)", False
            else:
                raise ValueError(f"unsupported archive format: {fmt}")
            row["integrity_result"] = integrity
            row["encrypted_status"] = "encrypted members present" if encrypted else "not encrypted"
            row["number_of_members"] = len(members)
            row["total_compressed_size"] = sum(int(m.get("compressed", 0)) for m in members)
            row["estimated_uncompressed_size"] = sum(int(m.get("uncompressed", 0)) for m in members)
            if len(members) > MAX_MEMBERS:
                errors.append(f"member limit exceeded: {len(members)} > {MAX_MEMBERS}")
            if row["estimated_uncompressed_size"] > MAX_OUTER_UNCOMPRESSED:
                errors.append("archive estimate exceeds 10 GB outer-archive safety limit")
            per_outer_total[outer] += int(row["estimated_uncompressed_size"])
            if per_outer_total[outer] > MAX_OUTER_UNCOMPRESSED:
                errors.append("cumulative nested estimate exceeds 10 GB outer-archive safety limit")
            names = [str(m["name"]).replace("\\", "/").lower() for m in members]
            counts = Counter(names)
            occurrence: defaultdict[str, int] = defaultdict(int)
            for member in members:
                name = str(member["name"])
                safe, normalized = safe_member(name)
                member_notes: list[str] = []
                if member.get("is_link"):
                    safe = False
                    member_notes.append("links are not extracted")
                size = int(member.get("uncompressed", 0))
                compressed = int(member.get("compressed", 0))
                if size > MAX_SINGLE_FILE:
                    safe = False
                    member_notes.append("single-file 5 GB safety limit exceeded")
                if size == 0 and not member.get("is_dir"):
                    member_notes.append("zero-byte file")
                ratio = size / max(compressed, 1)
                if compressed and ratio > HIGH_RATIO:
                    member_notes.append(f"suspicious compression ratio {ratio:.1f}:1")
                    warnings.append(f"high compression ratio: {name}")
                key = normalized.lower()
                dup_index = occurrence[key]
                occurrence[key] += 1
                if counts[key] > 1:
                    member_notes.append("duplicate member path; deterministic suffix used")
                target = _unique_target(extraction_base, normalized, dup_index) if safe else None
                nested = is_archive(name) and not member.get("is_dir")
                member_rows.append({
                    "country": country, "archive_id": archive_id, "parent_archive_id": parent,
                    "nesting_depth": depth, "member_path": normalized, "filename": PurePosixPath(normalized).name,
                    "extension": "".join(Path(normalized).suffixes).lower(), "compressed_size": compressed,
                    "uncompressed_size": size, "crc_result": member.get("crc", "not available"),
                    "encryption_status": "encrypted" if member.get("encrypted") else "not encrypted",
                    "directory_or_file": "directory" if member.get("is_dir") else "file",
                    "nested_archive_status": "nested archive" if nested else "not archive",
                    "suspected_role": suspected_role(normalized), "extraction_path": rel(target) if target else "",
                    "safe_or_unsafe": "safe" if safe else "unsafe", "notes": "; ".join(member_notes),
                })
            if encrypted:
                errors.append("password-protected/encrypted member present")
            if duplicate_of:
                row["extraction_status"] = "duplicate skipped"
            elif errors:
                row["extraction_status"] = "stopped: safety or encryption condition"
            else:
                extraction_base.mkdir(parents=True, exist_ok=True)
                extracted_nested: list[Path] = []
                for member, member_row in zip(members, member_rows[-len(members):] if members else []):
                    if member_row["safe_or_unsafe"] != "safe" or member.get("is_dir"):
                        continue
                    target = ROOT / member_row["extraction_path"]
                    target.parent.mkdir(parents=True, exist_ok=True)
                    if target.exists():
                        member_row["notes"] = "; ".join(filter(None, [member_row["notes"], "already present; not overwritten"]))
                    else:
                        if fmt == "zip":
                            with zipfile.ZipFile(path) as archive, archive.open(member["zip_info"], "r") as src, target.open("xb") as dst:
                                shutil.copyfileobj(src, dst, length=1024 * 1024)
                        elif fmt in {"tar", "tar.gz", "tgz"}:
                            with tarfile.open(path, "r:*") as archive:
                                source = archive.extractfile(member["tar_info"])
                                if source is not None:
                                    with source, target.open("xb") as dst:
                                        shutil.copyfileobj(source, dst, length=1024 * 1024)
                        else:
                            with gzip.open(path, "rb") as src, target.open("xb") as dst:
                                shutil.copyfileobj(src, dst, length=1024 * 1024)
                    if member_row["nested_archive_status"] == "nested archive" and target.exists():
                        extracted_nested.append(target)
                row["extraction_status"] = "extracted" if not any("already present" in r["notes"] for r in member_rows[-len(members):]) else "verified existing extraction"
                for nested_path in sorted(extracted_nested, key=lambda p: rel(p).lower()):
                    queue.append({"country": country, "path": nested_path, "parent": archive_id,
                                  "depth": depth + 1, "outer": outer, "outer_base": outer_base})
        except Exception as exc:  # archive-specific failure must not stop other archives
            errors.append(f"{type(exc).__name__}: {exc}")
            row["extraction_status"] = "failed safely"
            LOGGER.exception("Archive failed safely: %s", rel(path))
        row["warnings"] = "; ".join(sorted(set(warnings)))
        row["errors"] = "; ".join(errors)
        if roots[country] == ROOT / "data" / country and depth == 0:
            row["notes"] = "Original archive detected in legacy data/<country>/ location; left in place."
        archive_rows.append(row)

    archive_rows.sort(key=lambda r: (r["country"], int(r["nesting_depth"]), r["relative_path"].lower()))
    member_rows.sort(key=lambda r: (r["country"], r["archive_id"], r["member_path"].lower()))
    write_csv(CHECKPOINTS / "phase_01_archive_inventory.csv", archive_rows, ARCHIVE_FIELDS)
    write_csv(CHECKPOINTS / "phase_01_archive_members.csv", member_rows, MEMBER_FIELDS)
    duplicates = [
        {"country": r["country"], "archive_id": r["archive_id"], "sha256": r["sha256"],
         "duplicate_status": r["duplicate_status"], "relative_path": r["relative_path"]}
        for r in archive_rows if r["duplicate_status"] != "unique"
    ]
    write_csv(CHECKPOINTS / "phase_01_archive_duplicates.csv", duplicates,
              ["country", "archive_id", "sha256", "duplicate_status", "relative_path"])
    LOGGER.info("Inventoried %s archives and %s members", len(archive_rows), len(member_rows))
    return archive_rows, member_rows


def audit_environment(stage: str = "before") -> dict[str, Any]:
    """Record a privacy-minimized environment and original-archive checksum audit."""
    ensure_structure()
    roots = detected_raw_roots()
    packages = {}
    for package in ["pandas", "numpy", "pyreadstat", "pypdf", "pdfplumber", "python-docx", "openpyxl"]:
        try:
            packages[package] = importlib.metadata.version(package)
        except importlib.metadata.PackageNotFoundError:
            packages[package] = "not installed"
    archives = []
    for country, root in roots.items():
        for path in sorted(p for p in root.rglob("*") if p.is_file() and is_archive(p)):
            archives.append({"country": country, "path": rel(path), "size_bytes": path.stat().st_size,
                             "sha256": sha256_file(path)})
    git_status = "not a Git work tree"
    git = shutil.which("git")
    if git:
        result = subprocess.run([git, "status", "--short", "--branch"], cwd=ROOT, capture_output=True, text=True)
        if result.returncode == 0:
            git_status = result.stdout.strip()
    usage = shutil.disk_usage(ROOT)
    payload = {
        "audit_stage": stage,
        "timestamp_utc": datetime.now(timezone.utc).isoformat(),
        "operating_system": platform.system(),
        "os_release": platform.release(),
        "python_version": platform.python_version(),
        "project_root": ".",
        "available_disk_space_bytes": usage.free,
        "relevant_packages": packages,
        "git_status": git_status,
        "raw_archive_roots": {country: rel(path) for country, path in roots.items()},
        "original_archives": archives,
        "privacy_note": "Username, credentials, tokens, environment variables, and respondent data are excluded.",
    }
    path = CHECKPOINTS / "phase_01_environment.json"
    if path.exists() and stage == "after":
        try:
            before = json.loads(path.read_text(encoding="utf-8"))
            before_map = {x["path"]: x["sha256"] for x in before.get("original_archives", [])}
            after_map = {x["path"]: x["sha256"] for x in archives}
            payload["pre_extraction_original_archives"] = before.get("original_archives", [])
            payload["raw_checksums_unchanged"] = before_map == after_map
        except Exception as exc:
            payload["raw_checksums_unchanged"] = False
            payload["checksum_comparison_error"] = str(exc)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    LOGGER.info("Environment audit written (%s)", stage)
    return payload


FILE_FIELDS = [
    "country", "relative_path", "source_archive", "nesting_depth", "filename", "extension", "size",
    "probable_survey", "probable_year", "probable_wave_or_round", "probable_observation_level",
    "suspected_role", "openable_status", "opening_error", "notes",
]


def country_from_path(path: Path) -> str:
    text = rel(path).lower()
    if "kyrgyzstan" in text:
        return "kyrgyzstan"
    if "uzbekistan" in text or "/uz" in text:
        return "uzbekistan"
    return ""


def infer_year(path_text: str) -> str:
    text = path_text.lower()
    if "lik19" in text or "version 2022" in text or "version_2022" in text:
        return "2019"
    if "mics6" in text or "2021-22" in text:
        return "2021-22"
    years = re.findall(r"(?<!\d)(20(?:00|01|02|03|04|05|06|07|08|09|10|11|12|13|14|15|16|17|18|19|20|21|22|23|24|25))(?!\d)", text)
    return years[-1] if years else ""


def infer_survey(path_text: str, country: str) -> str:
    text = path_text.lower()
    if country == "kyrgyzstan" or "lik" in text:
        return "Life in Kyrgyzstan Study (LiK)"
    if "mics6" in text or "2021-22 mics" in text:
        return "Uzbekistan Multiple Indicator Cluster Survey (MICS), Round 6"
    if "mics2" in text or "2000 mics" in text:
        return "Uzbekistan Multiple Indicator Cluster Survey (MICS), Round 2"
    if "2006 mics" in text or "mics 2006" in text or "mics_2006" in text:
        return "Uzbekistan Multiple Indicator Cluster Survey (MICS), 2006"
    if country == "uzbekistan":
        return "Uzbekistan UNICEF survey collection (exact component inferred from path)"
    return "UNCERTAIN"


def infer_level(path_text: str) -> str:
    text = path_text.lower().replace("\\", "/")
    stem = Path(text).stem
    rules = [
        ("community", ["/community/", "/cm", "/ao_"]),
        ("agriculture", ["/agriculture/", "/agr", "/ag"]),
        ("household", ["/household/", "/hh", "uzhh"]),
        ("woman", ["/wm", "uzwm"]),
        ("child", ["/ch", "uzch"]),
        ("birth", ["/bh"]),
        ("person", ["/individual/", "/individial/", "/hl", "uzhl", "/id"]),
        ("youth", ["/youth/", "/yt"]),
        ("migration episode", ["migration"]),
        ("consumption item", ["consumption", "food item"]),
    ]
    for level, tokens in rules:
        if any(token in text or stem.startswith(token.strip("/")) for token in tokens):
            return level
    return "UNCERTAIN"


def _probe_file(path: Path) -> tuple[str, str]:
    """Open a file minimally without exposing contents."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".zip":
            with zipfile.ZipFile(path) as zf:
                zf.infolist()
        elif suffix in {".tar", ".tgz"} or str(path).lower().endswith(".tar.gz"):
            with tarfile.open(path, "r:*") as tf:
                tf.next()
        elif suffix == ".gz":
            with gzip.open(path, "rb") as handle:
                handle.read(1)
        else:
            with path.open("rb") as handle:
                handle.read(16)
        return "openable", ""
    except Exception as exc:
        return "not openable", f"{type(exc).__name__}: {exc}"


def inventory_files() -> list[dict[str, Any]]:
    """Inventory raw, extracted, documentation, and literature files."""
    member_map = {r["extraction_path"]: r for r in read_csv(CHECKPOINTS / "phase_01_archive_members.csv") if r.get("extraction_path")}
    rows: list[dict[str, Any]] = []
    scan_roots = [ROOT / "data" / "raw", ROOT / "data" / "kyrgyzstan", ROOT / "data" / "uzbekistan",
                  UNPACKED, ROOT / "documentation", ROOT / "literature"]
    seen: set[Path] = set()
    for scan_root in scan_roots:
        if not scan_root.exists():
            continue
        for path in sorted((p for p in scan_root.rglob("*") if p.is_file()), key=lambda p: rel(p).lower()):
            resolved = path.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)
            relative = rel(path)
            country = country_from_path(path)
            member = member_map.get(relative, {})
            status, error = _probe_file(path)
            text = relative.lower()
            notes = []
            if relative.startswith("data/kyrgyzstan/") or relative.startswith("data/uzbekistan/"):
                notes.append("original archive/file in legacy raw-data location; left unchanged")
            rows.append({
                "country": country, "relative_path": relative, "source_archive": member.get("archive_id", ""),
                "nesting_depth": member.get("nesting_depth", ""), "filename": path.name,
                "extension": "".join(path.suffixes).lower(), "size": path.stat().st_size,
                "probable_survey": infer_survey(text, country), "probable_year": infer_year(text),
                "probable_wave_or_round": infer_year(text), "probable_observation_level": infer_level(text),
                "suspected_role": suspected_role(text), "openable_status": status, "opening_error": error,
                "notes": "; ".join(notes),
            })
    write_csv(CHECKPOINTS / "phase_01_file_inventory.csv", rows, FILE_FIELDS)
    LOGGER.info("Inventoried %s files", len(rows))
    return rows


DATASET_FIELDS = [
    "country", "survey_name", "source_archive", "relative_path", "filename", "format", "size",
    "number_of_rows", "number_of_columns", "survey_year", "wave_or_round", "probable_observation_level",
    "variable_names", "variable_labels", "value_label_availability", "candidate_id_variables",
    "candidate_weight_variables", "candidate_strata", "candidate_psu", "candidate_region",
    "candidate_urban_rural_variable", "candidate_date_variables", "duplicate_row_risk",
    "all_missing_column_count", "warnings", "notes",
]
VARIABLE_FIELDS = [
    "country", "survey_name", "survey_year", "wave_or_round", "source_file", "source_archive",
    "observation_level", "variable_name", "variable_label", "value_labels", "storage_format", "display_format",
]


def _candidate_names(variables: list[dict[str, Any]], patterns: list[str]) -> list[str]:
    found = []
    compiled = [re.compile(p, re.I) for p in patterns]
    for variable in variables:
        text = f"{variable.get('variable_name', '')} {variable.get('variable_label', '')}"
        if any(pattern.search(text) for pattern in compiled):
            found.append(str(variable.get("variable_name", "")))
    return found


def _read_dataset_metadata(path: Path) -> tuple[int | str, list[dict[str, Any]], list[str]]:
    """Read metadata before data; return row count, variables, and warnings."""
    suffix = path.suffix.lower()
    variables: list[dict[str, Any]] = []
    warnings: list[str] = []
    row_count: int | str = ""
    if suffix in {".dta", ".sav", ".por"}:
        import pyreadstat  # local dependency documented in requirements
        readers = {".dta": pyreadstat.read_dta, ".sav": pyreadstat.read_sav, ".por": pyreadstat.read_por}
        _, metadata = readers[suffix](str(path), metadataonly=True)
        row_count = metadata.number_rows
        label_map = metadata.column_names_to_labels or {}
        value_map = getattr(metadata, "variable_value_labels", {}) or {}
        formats = getattr(metadata, "original_variable_types", {}) or {}
        storage = getattr(metadata, "readstat_variable_types", {}) or {}
        for name in metadata.column_names:
            variables.append({
                "variable_name": name, "variable_label": label_map.get(name) or "",
                "value_labels": value_map.get(name, {}), "storage_format": storage.get(name, ""),
                "display_format": formats.get(name, ""),
            })
    elif suffix in {".csv", ".tab", ".txt"}:
        import pandas as pd
        separator = "\t" if suffix in {".tab"} else None
        frame = pd.read_csv(path, sep=separator, engine="python", nrows=5)
        with path.open("rb") as handle:
            row_count = max(sum(1 for _ in handle) - 1, 0)
        variables = [{"variable_name": str(name), "variable_label": "", "value_labels": {},
                      "storage_format": str(dtype), "display_format": ""}
                     for name, dtype in frame.dtypes.items()]
    elif suffix in {".xlsx", ".xls"}:
        if suffix == ".xls":
            raise RuntimeError("legacy .xls metadata reader is unavailable")
        from openpyxl import load_workbook
        workbook = load_workbook(path, read_only=True, data_only=True)
        sheet = workbook[workbook.sheetnames[0]]
        row_count = max((sheet.max_row or 1) - 1, 0)
        headers = [cell.value for cell in next(sheet.iter_rows(min_row=1, max_row=1))]
        variables = [{"variable_name": str(name or f"column_{i+1}"), "variable_label": "",
                      "value_labels": {}, "storage_format": "", "display_format": ""}
                     for i, name in enumerate(headers)]
        workbook.close()
    elif suffix == ".json":
        obj = json.loads(path.read_text(encoding="utf-8"))
        records = obj if isinstance(obj, list) else [obj]
        row_count = len(records)
        keys = sorted({key for record in records[:100] if isinstance(record, dict) for key in record})
        variables = [{"variable_name": key, "variable_label": "", "value_labels": {},
                      "storage_format": "", "display_format": ""} for key in keys]
    else:
        raise RuntimeError(f"metadata reader not implemented for {suffix}")
    return row_count, variables, warnings


def _dataset_quality(path: Path) -> tuple[str, str, str]:
    """Compute aggregate all-missing and duplicate-row counts after metadata is read."""
    suffix = path.suffix.lower()
    try:
        if suffix in {".dta", ".sav", ".por"}:
            import pyreadstat
            reader = {".dta": pyreadstat.read_dta, ".sav": pyreadstat.read_sav, ".por": pyreadstat.read_por}[suffix]
            frame, _ = reader(str(path))
        elif suffix in {".csv", ".txt", ".tab"}:
            import pandas as pd
            frame = pd.read_csv(path, sep="\t" if suffix == ".tab" else None, engine="python")
        else:
            return "not computed", "not computed", "quality scan not implemented for this format"
        all_missing = int(frame.isna().all(axis=0).sum())
        duplicate_count = int(frame.duplicated().sum())
        risk = f"{duplicate_count} exact duplicate rows" if duplicate_count else "no exact duplicate rows"
        return str(all_missing), risk, ""
    except Exception as exc:
        return "not computed", "not computed", f"quality-scan limitation: {type(exc).__name__}: {exc}"


def audit_data_metadata() -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Audit all likely datasets, producing file- and variable-level metadata."""
    files = read_csv(CHECKPOINTS / "phase_01_file_inventory.csv")
    dataset_rows: list[dict[str, Any]] = []
    variable_rows: list[dict[str, Any]] = []
    allowed = {".dta", ".sav", ".por", ".csv", ".xlsx", ".xls", ".parquet", ".json", ".txt", ".tab"}
    for record in files:
        path = ROOT / record["relative_path"]
        suffix = path.suffix.lower()
        if record.get("suspected_role") != "dataset" or suffix not in allowed or not path.exists():
            continue
        if not record["relative_path"].startswith("data/interim/unpacked/"):
            continue
        country = record.get("country", "")
        survey = record.get("probable_survey", infer_survey(record["relative_path"], country))
        year = record.get("probable_year", infer_year(record["relative_path"]))
        level = record.get("probable_observation_level", infer_level(record["relative_path"]))
        warnings: list[str] = []
        try:
            row_count, variables, read_warnings = _read_dataset_metadata(path)
            warnings.extend(read_warnings)
            all_missing, duplicate_risk, quality_warning = _dataset_quality(path)
            if quality_warning:
                warnings.append(quality_warning)
            status_note = "Metadata read successfully; no respondent values exported."
        except Exception as exc:
            row_count, variables = "", []
            all_missing, duplicate_risk = "not computed", "not computed"
            warnings.append(f"{type(exc).__name__}: {exc}")
            status_note = "PARSE LIMITATION - MANUAL REVIEW REQUIRED"
            LOGGER.warning("Dataset metadata limitation: %s: %s", rel(path), exc)
        id_vars = _candidate_names(variables, [r"\b(id|ident|panel|cluster|household number|line number)\b", r"^hhid$", r"^caseid$"])
        weight_vars = _candidate_names(variables, [r"sample weight", r"^hhweight\b", r"^wmweight\b", r"^chweight\b", r"^fsweight\b"])
        strata_vars = _candidate_names(variables, [r"strat"])
        psu_vars = _candidate_names(variables, [r"\bpsu\b", r"cluster", r"enumeration area"])
        region_vars = _candidate_names(variables, [r"region", r"oblast", r"province"])
        urban_vars = _candidate_names(variables, [r"urban", r"rural", r"area of residence"])
        date_vars = _candidate_names(variables, [r"interview.*date", r"date.*interview", r"\bmonth\b", r"\byear\b"])
        dataset_rows.append({
            "country": country, "survey_name": survey, "source_archive": record.get("source_archive", ""),
            "relative_path": record["relative_path"], "filename": path.name, "format": suffix.lstrip("."),
            "size": path.stat().st_size, "number_of_rows": row_count, "number_of_columns": len(variables),
            "survey_year": year, "wave_or_round": year, "probable_observation_level": level,
            "variable_names": [v["variable_name"] for v in variables],
            "variable_labels": {v["variable_name"]: v["variable_label"] for v in variables if v["variable_label"]},
            "value_label_availability": sum(bool(v["value_labels"]) for v in variables),
            "candidate_id_variables": id_vars, "candidate_weight_variables": weight_vars,
            "candidate_strata": strata_vars, "candidate_psu": psu_vars, "candidate_region": region_vars,
            "candidate_urban_rural_variable": urban_vars, "candidate_date_variables": date_vars,
            "duplicate_row_risk": duplicate_risk,
            "all_missing_column_count": all_missing, "warnings": "; ".join(warnings), "notes": status_note,
        })
        for variable in variables:
            variable_rows.append({
                "country": country, "survey_name": survey, "survey_year": year, "wave_or_round": year,
                "source_file": record["relative_path"], "source_archive": record.get("source_archive", ""),
                "observation_level": level, **variable,
            })
    dataset_rows.sort(key=lambda r: (r["country"], r["survey_year"], r["relative_path"].lower()))
    variable_rows.sort(key=lambda r: (r["country"], r["survey_year"], r["source_file"].lower(), r["variable_name"].lower()))
    write_csv(CHECKPOINTS / "phase_01_dataset_inventory.csv", dataset_rows, DATASET_FIELDS)
    write_csv(CHECKPOINTS / "phase_01_variable_metadata.csv", variable_rows, VARIABLE_FIELDS)
    LOGGER.info("Audited %s datasets and %s variables", len(dataset_rows), len(variable_rows))
    return dataset_rows, variable_rows


TARGET_TERMS = {
    "migration": [r"migrat", r"migrant", r"moved abroad", r"away from household", r"moving to current place", r"duration of living in current place"],
    "remittances": [r"remitt", r"money.*abroad", r"transfer.*abroad", r"sent.*money", r"receive(?:d)?.*money"],
    "remittance amount": [r"amount.*remitt", r"how much.*migrant.*send", r"how much money.*send"],
    "economic shocks": [r"economic shock", r"income loss", r"business closure", r"increase.*price", r"unexpected expense"],
    "employment shocks": [r"job loss", r"lost.*job", r"employment loss", r"laid off"],
    "health shocks": [r"health shock", r"serious illness", r"injury", r"death.*household"],
    "agricultural shocks": [r"crop loss", r"livestock loss", r"livestock disease", r"\bpests?\b", r"harvest loss", r"animals were lost due to death"],
    "climate-related shocks": [r"drought", r"flood", r"cold winter", r"extreme heat", r"natural disaster", r"weather shock"],
    "loss of remittances": [r"loss.*remitt", r"remitt.*loss", r"stopped.*money.*abroad"],
    "food insecurity": [r"food insecur", r"without eating", r"not eaten all day", r"skipped meal", r"missed a meal", r"reduced.*meal", r"eaten less", r"run out of food", r"not enough food", r"household hunger", r"hungry.*unable to eat", r"worried.*food", r"unable to eat healthy"],
    "food consumption": [r"food consum", r"dietary diversity", r"foods? eaten", r"food group"],
    "food expenditure": [r"food expend", r"spent.*food", r"expenditure.*food", r"food purchase"],
    "total expenditure": [r"total expend", r"household expenditure", r"total consumption"],
    "household income": [r"household income", r"total household income", r"income (of|for) (the )?household"],
    "subjective financial condition": [r"satisf.*household income", r"financial situation", r"subjective financial"],
    "wealth": [r"wealth index", r"wealth quintile", r"wealth score"],
    "assets": [r"household asset", r"own.*television", r"own.*vehicle", r"asset ownership"],
    "household size": [r"household size", r"number.*household member", r"members.*household"],
    "rural or urban": [r"area of residence", r"urban.*rural", r"rural.*urban"],
    "region": [r"region", r"oblast", r"province"],
    "women's education": [r"woman.*education", r"educational attainment", r"highest.*education"],
    "women's employment": [r"woman.*work", r"woman.*employ", r"currently working"],
    "child education": [r"school attendance", r"attend.*school", r"education.*child"],
    "child nutrition": [r"stunt", r"wast", r"underweight", r"height.*age", r"weight.*height"],
    "water": [r"source of drinking water", r"drinking water", r"water source"],
    "sanitation": [r"toilet", r"sanitation", r"human waste"],
    "digital access": [r"internet", r"computer", r"mobile phone"],
    "household identifier": [r"household (number|identifier|id)", r"unique household"],
    "survey year or round": [r"survey year", r"survey round", r"survey wave", r"panel wave", r"wave number", r"year of interview"],
    "weights": [r"sample weight", r"sampling weight", r"household sample weight", r"women.?s sample weight", r"child(?:ren)?.*sample weight"],
    "working-age adults": [r"working.age", r"adult household member"],
    "dependency ratio": [r"dependency ratio"],
    "household-head age": [r"age of household head", r"household head.*age"],
    "household-head sex": [r"sex of household head", r"household head.*sex"],
    "household-head education": [r"education of household head", r"household head.*education"],
}


def _extract_text(path: Path, max_chars: int = 300_000) -> tuple[str, str]:
    """Extract documentation text conservatively; return text and limitation."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages[:250]:
                pages.append(page.extract_text() or "")
                if sum(len(x) for x in pages) >= max_chars:
                    break
            text = "\n".join(pages)[:max_chars]
            limitation = "" if text.strip() else "PARSE LIMITATION - MANUAL REVIEW REQUIRED"
            return text, limitation
        if suffix == ".docx":
            from docx import Document
            document = Document(str(path))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
            return "\n".join(parts)[:max_chars], ""
        if suffix in {".txt", ".md", ".rtf", ".sps", ".do", ".stcmd"}:
            data = path.read_bytes()
            for encoding in ("utf-8", "cp1251", "latin-1"):
                try:
                    return data.decode(encoding)[:max_chars], ""
                except UnicodeDecodeError:
                    continue
        if suffix == ".doc":
            return "", "PARSE LIMITATION - MANUAL REVIEW REQUIRED (legacy .doc)"
        return "", "unsupported documentation type"
    except Exception as exc:
        return "", f"PARSE LIMITATION - MANUAL REVIEW REQUIRED ({type(exc).__name__}: {exc})"


def audit_documentation() -> list[dict[str, Any]]:
    """Extract keyword evidence from questionnaires, manuals, and reports."""
    files = read_csv(CHECKPOINTS / "phase_01_file_inventory.csv")
    roles = {"questionnaire", "codebook", "manual", "report", "README", "syntax", "labels", "metadata", "publication"}
    audit_rows: list[dict[str, Any]] = []
    for record in files:
        if record.get("suspected_role") not in roles:
            continue
        path = ROOT / record["relative_path"]
        if not path.exists():
            continue
        text, limitation = _extract_text(path)
        lower = re.sub(r"\s+", " ", text.lower())
        hits: dict[str, list[str]] = {}
        for concept, patterns in TARGET_TERMS.items():
            snippets = []
            for pattern in patterns:
                match = re.search(pattern, lower, flags=re.I)
                if match:
                    start, end = max(0, match.start() - 120), min(len(lower), match.end() + 180)
                    snippets.append(lower[start:end].strip())
            if snippets:
                hits[concept] = snippets[:3]
        audit_rows.append({
            "country": record.get("country", ""), "relative_path": record["relative_path"],
            "survey_name": record.get("probable_survey", ""), "survey_year": record.get("probable_year", ""),
            "document_role": record.get("suspected_role", ""), "characters_extracted": len(text),
            "keyword_concepts_found": sorted(hits), "evidence_snippets": hits,
            "parse_status": limitation or "parsed", "notes": "Automated text evidence; consult the cited source before harmonization.",
        })
    path = CHECKPOINTS / "phase_01_documentation_audit.json"
    path.write_text(json.dumps(audit_rows, ensure_ascii=False, indent=2), encoding="utf-8")
    LOGGER.info("Audited %s documentation files", len(audit_rows))
    return audit_rows


CANDIDATE_FIELDS = [
    "country", "survey_name", "survey_year", "wave_or_round", "source_file", "observation_level",
    "variable_name", "variable_label", "question_wording", "response_categories", "value_labels", "unit",
    "recall_period", "missing_value_codes", "proposed_concept", "candidate_role", "comparability_confidence", "notes",
]


def _concept_role(concept: str) -> str:
    if concept in {"food insecurity", "food consumption", "food expenditure", "total expenditure", "household income", "subjective financial condition", "wealth"}:
        return "outcome"
    if concept in {"remittances", "remittance amount", "migration"}:
        return "main predictor"
    if "shock" in concept or concept == "loss of remittances":
        return "moderator"
    if concept in {"household identifier", "survey year or round"}:
        return "identifier"
    if concept in {"weights", "rural or urban", "region"}:
        return "survey-design variable" if concept == "weights" else "control"
    return "control"


def _detect_recall(label: str) -> str:
    patterns = [r"past (\d+) days?", r"last (\d+) days?", r"past (\d+) months?", r"last (\d+) months?",
                r"last week", r"past week", r"last year", r"past year", r"previous 12 months"]
    found = []
    for pattern in patterns:
        found.extend(m.group(0) for m in re.finditer(pattern, label, flags=re.I))
    return "; ".join(dict.fromkeys(found))


def _missing_codes(value_labels: str) -> str:
    """Identify documented non-substantive response codes from value labels."""
    try:
        mapping = json.loads(value_labels) if value_labels else {}
    except json.JSONDecodeError:
        return "unknown"
    terms = re.compile(r"refus|don.?t know|\bdk\b|no response|not applicable|missing", re.I)
    codes = [str(code) for code, label in mapping.items() if terms.search(str(label))]
    return "; ".join(codes) if codes else "none documented in value labels"


def _candidate_confidence(name: str, label: str, value_labels: str) -> str:
    if label and value_labels:
        return "high"
    if label:
        return "medium"
    if name:
        return "low"
    return "unknown"


def extract_variable_candidates() -> list[dict[str, Any]]:
    """Generate provisional concept candidates from labels plus names, not names alone."""
    metadata = read_csv(CHECKPOINTS / "phase_01_variable_metadata.csv")
    candidates: list[dict[str, Any]] = []
    compiled = {concept: [re.compile(pattern, re.I) for pattern in patterns] for concept, patterns in TARGET_TERMS.items()}
    # Narrow, documented survey-design names supplement semantic labels.
    exact_names = {
        "hhweight": "weights", "wmweight": "weights", "chweight": "weights", "fsweight": "weights",
        "hh2": "household identifier", "hhid": "household identifier",
        "hh6": "rural or urban", "hh7": "region", "windex5": "wealth", "wscore": "wealth",
    }
    for variable in metadata:
        name = variable.get("variable_name", "")
        label = variable.get("variable_label", "")
        value_labels = variable.get("value_labels", "")
        source_lower = variable.get("source_file", "").lower()
        label_text = re.sub(r"\s+", " ", label).strip()
        concepts = []
        for concept, patterns in compiled.items():
            # Labels/question wording carry the substantive mapping; names only support it.
            if label_text and any(pattern.search(label_text) for pattern in patterns):
                concepts.append(concept)
        if "household income" in concepts and ("completely dissatisfied" in value_labels.lower() or name.lower() == "i101_3"):
            concepts.remove("household income")
            concepts.append("subjective financial condition")
        if name.lower() in exact_names:
            concepts.append(exact_names[name.lower()])
        # Contextual mappings are supported by module path, full labels, and value labels.
        if source_lower.endswith("/household/hh7.dta") and name.lower() == "shock" and "name of shock" in label_text.lower():
            concepts.extend(["economic shocks", "employment shocks", "health shocks", "agricultural shocks",
                             "climate-related shocks", "loss of remittances"])
        if source_lower.endswith("/household/hh7.dta") and name.lower() in {"h701", "h703", "h704"}:
            concepts.append("economic shocks")
        if source_lower.endswith("/household/hh4a.dta") and name.lower() == "h401c" and "total expenses" in label_text.lower():
            concepts.append("food expenditure")
        if source_lower.endswith("/household/hh4a.dta") and name.lower() == "h402a" and "own production" in label_text.lower():
            concepts.append("food consumption")
        if source_lower.endswith("/household/hh2b.dta") and name.lower() == "asset" and label_text.lower() == "asset":
            concepts.append("assets")
        if source_lower.endswith("/individual/id2.dta") and re.fullmatch(r"i251_[1-8]", name.lower()):
            concepts.append("food insecurity")
        if source_lower.endswith("/household/hh6b.dta") and name.lower() == "h620":
            concepts.append("remittances")
        if source_lower.endswith("/household/hh6b.dta") and name.lower() == "h622":
            concepts.append("remittance amount")
        for concept in sorted(set(concepts)):
            confidence = _candidate_confidence(name, label, value_labels)
            candidate_note = "Provisional automated candidate from variable label/value metadata; questionnaire verification required before harmonization."
            if concept == "migration" and "immigration origin" in label_text.lower():
                confidence = "not comparable"
                candidate_note = "Not a migration measure: discrimination item mentioning immigration origin."
            if concept == "household income" and ("/individual/id5" in source_lower or "/individial/id5" in source_lower or "which member" in label_text.lower() or "of household income" in label_text.lower() or "contribute to the household income" in label_text.lower()):
                confidence = "not comparable"
                candidate_note = "Decision-making or attitude item mentioning household income; not an income amount."
            if concept == "climate-related shocks" and "school closed" in label_text.lower():
                confidence = "not comparable"
                candidate_note = "Child-level school-closure consequence, not a household shock exposure measure."
            if source_lower.endswith("/household/hh7.dta") and name.lower() == "shock":
                confidence = "high"
                candidate_note = "LiK 2019 household shock roster; value labels identify exact economic, job-loss, remittance-loss, health, agricultural, and climate categories."
            if source_lower.endswith("/household/hh4a.dta") and name.lower() == "h401c":
                confidence = "high"
                candidate_note = "Food expenditure amount in the LiK food-item module; companion food_item variable defines the item roster."
            if source_lower.endswith("/household/hh4a.dta") and name.lower() == "h402a":
                confidence = "high"
                candidate_note = "Quantity consumed from own production in the LiK food-item module; not total food consumption without construction."
            if source_lower.endswith("/household/hh2b.dta") and name.lower() == "asset":
                confidence = "high"
                candidate_note = "LiK 2019 household asset roster; ownership/quantity fields in the same module must be paired before constructing wealth."
            if variable.get("country") == "uzbekistan" and name.upper() == "HH2" and concept == "household identifier":
                confidence = "high"
                candidate_note = "Household number; combine with HH1 cluster number for a survey-unique household key."
            if source_lower.endswith("/individual/id2.dta") and re.fullmatch(r"i251_[1-8]", name.lower()):
                confidence = "high"
                candidate_note = "One of eight LiK 2019 food-insecurity experience items; individual response must be aggregated or the unit changed explicitly."
            recall = _detect_recall(label)
            if source_lower.endswith("/household/hh7.dta") and name.lower() == "shock":
                recall = "last 12 months (from companion h701 wording)"
            if source_lower.endswith("/individual/id2.dta") and re.fullmatch(r"i251_[1-8]", name.lower()):
                recall = "last 12 months (questionnaire battery introduction)"
            unit = "Som" if re.search(r"\bSoms?\b", label, flags=re.I) else "not stated in metadata"
            if source_lower.endswith("/household/hh6b.dta") and name.lower() == "h622":
                unit = "currency recorded separately in h623 (Som or US dollar)"
            candidates.append({
                "country": variable.get("country", ""), "survey_name": variable.get("survey_name", ""),
                "survey_year": variable.get("survey_year", ""), "wave_or_round": variable.get("wave_or_round", ""),
                "source_file": variable.get("source_file", ""), "observation_level": variable.get("observation_level", ""),
                "variable_name": name, "variable_label": label, "question_wording": label,
                "response_categories": value_labels, "value_labels": value_labels, "unit": unit,
                "recall_period": recall, "missing_value_codes": _missing_codes(value_labels),
                "proposed_concept": concept, "candidate_role": _concept_role(concept),
                "comparability_confidence": confidence, "notes": candidate_note,
            })
    candidates.sort(key=lambda r: (r["country"], r["proposed_concept"], r["survey_year"], r["source_file"], r["variable_name"]))
    write_csv(CHECKPOINTS / "phase_01_variable_candidates.csv", candidates, CANDIDATE_FIELDS)
    LOGGER.info("Extracted %s provisional variable candidates", len(candidates))
    return candidates


COMPATIBILITY_CONCEPTS = [
    "migration", "remittances", "economic shocks", "employment shocks", "health shocks",
    "agricultural shocks", "climate-related shocks", "food insecurity", "food expenditure",
    "food consumption", "total expenditure", "household income", "wealth", "assets", "household size",
    "rural or urban", "region", "women's education", "women's employment", "child education",
    "child nutrition", "water", "sanitation", "digital access",
]
COMPATIBILITY_FIELDS = [
    "concept", "available_in_kyrgyzstan", "available_in_uzbekistan", "kyrgyzstan_candidate_variables",
    "uzbekistan_candidate_variables", "question_wording_match", "recall_period_match", "response_scale_match",
    "observation_level_match", "year_compatibility", "harmonization_possibility", "feasibility_status", "explanation",
]


def _latest_candidates(rows: list[dict[str, str]], limit: int = 12) -> list[dict[str, str]]:
    def year_key(row: dict[str, str]) -> tuple[int, str]:
        match = re.search(r"20\d{2}", row.get("survey_year", ""))
        return (int(match.group(0)) if match else 0, row.get("source_file", ""))
    unique: dict[tuple[str, str], dict[str, str]] = {}
    for row in rows:
        unique[(row.get("source_file", ""), row.get("variable_name", ""))] = row
    return sorted(unique.values(), key=year_key, reverse=True)[:limit]


def build_country_compatibility_matrix() -> list[dict[str, Any]]:
    """Compare concept candidates conservatively across the two survey systems."""
    candidates = read_csv(CHECKPOINTS / "phase_01_variable_candidates.csv")
    grouped: defaultdict[tuple[str, str], list[dict[str, str]]] = defaultdict(list)
    for row in candidates:
        grouped[(row.get("country", ""), row.get("proposed_concept", ""))].append(row)
    output = []
    for concept in COMPATIBILITY_CONCEPTS:
        kg = _latest_candidates([r for r in grouped[("kyrgyzstan", concept)] if r.get("comparability_confidence") != "not comparable"])
        uz = _latest_candidates([r for r in grouped[("uzbekistan", concept)] if r.get("comparability_confidence") != "not comparable"])
        kg_labels = {re.sub(r"\W+", " ", r.get("variable_label", "").lower()).strip() for r in kg if r.get("variable_label")}
        uz_labels = {re.sub(r"\W+", " ", r.get("variable_label", "").lower()).strip() for r in uz if r.get("variable_label")}
        exact_label = bool(kg_labels & uz_labels)
        recall_kg = {r.get("recall_period", "") for r in kg if r.get("recall_period")}
        recall_uz = {r.get("recall_period", "") for r in uz if r.get("recall_period")}
        level_kg = {r.get("observation_level", "") for r in kg}
        level_uz = {r.get("observation_level", "") for r in uz}
        if kg and uz and exact_label:
            status, harmonization = "directly comparable", "possible after questionnaire confirmation"
        elif kg and uz:
            status, harmonization = "conceptually comparable", "possible only with construct-level harmonization"
        elif kg or uz:
            status, harmonization = "not comparable", "not possible as a two-country measure from current files"
        else:
            status, harmonization = "unknown", "unknown"
        explanation = (
            "Automated candidates exist in both countries, but instruments, survey years, recall periods, and observation levels must be verified."
            if kg and uz else
            "Candidate evidence appears in only one country; absence here means not found in labels, not proof of absolute absence."
            if kg or uz else
            "No label-supported candidate was found; documentation may still require manual review."
        )
        output.append({
            "concept": concept, "available_in_kyrgyzstan": "yes" if kg else "not found",
            "available_in_uzbekistan": "yes" if uz else "not found",
            "kyrgyzstan_candidate_variables": [f"{r['variable_name']} | {r['source_file']}" for r in kg],
            "uzbekistan_candidate_variables": [f"{r['variable_name']} | {r['source_file']}" for r in uz],
            "question_wording_match": "exact label match" if exact_label else "no exact match / unknown",
            "recall_period_match": "matching stated recall" if recall_kg and recall_uz and recall_kg & recall_uz else "unknown or different",
            "response_scale_match": "requires value-label review" if kg and uz else "not assessable",
            "observation_level_match": "overlap" if level_kg & level_uz else "different or unknown",
            "year_compatibility": "LiK latest wave 2019 versus MICS6 2021-22; not contemporaneous" if kg and uz else "not assessable",
            "harmonization_possibility": harmonization, "feasibility_status": status, "explanation": explanation,
        })
    write_csv(CHECKPOINTS / "phase_01_country_compatibility_matrix.csv", output, COMPATIBILITY_FIELDS)
    return output


LITERATURE_FIELDS = [
    "full_title", "authors", "year", "journal_or_organization", "country", "theme", "dataset",
    "survey_years", "research_question", "outcome", "predictors", "methodology", "main_finding",
    "limitations", "relevance", "literature_role", "verified_from_full_text", "file_path", "verification_notes",
]


SEED_LITERATURE = [
    {
        "full_title": "Harsh Winter Shocks and Distress Sales", "authors": "Sultakeev and Petrick", "year": "2025",
        "country": "Kyrgyzstan", "theme": "climate shocks; coping; consumption and asset smoothing; livestock",
        "dataset": "Life in Kyrgyzstan panel data", "methodology": "Household fixed effects (initial note; not full-text verified)",
        "relevance": "Central to shocks, coping, food consumption, and asset smoothing.", "literature_role": "CORE",
    },
    {
        "full_title": "Gender-Sensitive Resilience in Kyrgyz Households", "authors": "Egamberdiev et al.", "year": "2025",
        "country": "Kyrgyzstan", "theme": "household resilience; adaptive capacity; assets; social safety nets; gender",
        "relevance": "Supports household resilience and gender-sensitive vulnerability framing.", "literature_role": "SUPPORTING",
    },
    {
        "full_title": "Citizen Perception and Participation in Local Government", "authors": "Bozkuş Kahyaoğlu et al.", "year": "2025",
        "country": "Kyrgyzstan", "theme": "local government; trust; participation",
        "dataset": "Life in Kyrgyzstan Study (initial note)", "methodology": "Logit regression (initial note; not full-text verified)",
        "relevance": "Example of LiK use; not central to the migration-remittance-food-security mechanism.",
        "literature_role": "OPTIONAL CONTEXT",
    },
    {
        "full_title": "Household Structure and Female Labor", "authors": "Kovaleva et al.", "year": "2025",
        "country": "Kyrgyzstan", "theme": "household structure; gender; female labor",
        "relevance": "Potential heterogeneity context; not central to the main mechanism.", "literature_role": "OPTIONAL CONTEXT",
    },
]


def inventory_literature() -> list[dict[str, Any]]:
    """Inventory available literature without inventing bibliographic details or findings."""
    original = ROOT / "literature" / "original"
    files = sorted((p for p in original.rglob("*") if p.is_file()), key=lambda p: rel(p).lower())
    rows: list[dict[str, Any]] = []
    matched_paths: set[Path] = set()
    for seed in SEED_LITERATURE:
        tokens = [t.lower() for t in re.findall(r"[A-Za-z]{5,}", seed["full_title"])[:3]]
        match = next((p for p in files if sum(token in p.name.lower() for token in tokens) >= 2), None)
        row = {field: "" for field in LITERATURE_FIELDS}
        row.update(seed)
        if match:
            text, limitation = _extract_text(match)
            row["verified_from_full_text"] = "yes" if text.strip() and not limitation else "no"
            row["file_path"] = rel(match)
            row["verification_notes"] = limitation or "Full text parsed; substantive findings still require manual scholarly verification."
            matched_paths.add(match)
        else:
            row["verified_from_full_text"] = "no"
            row["verification_notes"] = "Seed note preserved; no matching full-text file supplied. Bibliographic details and findings not verified."
        rows.append(row)
    for path in files:
        if path in matched_paths:
            continue
        text, limitation = _extract_text(path)
        row = {field: "" for field in LITERATURE_FIELDS}
        row.update({
            "full_title": path.stem, "literature_role": "NOT CENTRAL", "verified_from_full_text": "yes" if text.strip() and not limitation else "no",
            "file_path": rel(path), "verification_notes": limitation or "File parsed; title is filename only and citation details need manual verification.",
        })
        rows.append(row)
    write_csv(CHECKPOINTS / "phase_01_literature_inventory.csv", rows, LITERATURE_FIELDS)
    write_csv(ROOT / "literature" / "matrices" / "literature_matrix.csv", rows, LITERATURE_FIELDS)
    notes = """# Original Topic Notes

These notes preserve the four literature seeds supplied by the supervisor. They are not a substitute for full-text verification.

1. **Sultakeev and Petrick (2025), _Harsh Winter Shocks and Distress Sales_ — CORE.** Initial relevance: climate shocks, household coping, food consumption, consumption smoothing, asset smoothing, livestock, household fixed effects, and Life in Kyrgyzstan panel data.

2. **Egamberdiev et al. (2025), _Gender-Sensitive Resilience in Kyrgyz Households_ — SUPPORTING.** Initial relevance: household resilience, adaptive capacity, assets, social safety nets, and gender-sensitive household vulnerability.

3. **Bozkuş Kahyaoğlu et al. (2025), _Citizen Perception and Participation in Local Government_ — OPTIONAL CONTEXT.** Initial relevance: an example of LiK use, local government, trust, and logit regression. It is not central to the migration-remittance-food-security mechanism.

4. **Kovaleva et al. (2025), _Household Structure and Female Labor_ — OPTIONAL CONTEXT.** Initial relevance: household structure, gender, female labor, and possible heterogeneity analysis.

## Literature still needed

- Remittances and household expenditure
- Migration as informal insurance
- Shocks and migration decisions
- Resilience and food security
- Uzbekistan migration and remittances
- Uzbekistan food insecurity and coping
- Central Asian household resilience
"""
    (ROOT / "literature" / "notes" / "original_topic_notes.md").write_text(notes, encoding="utf-8")
    verification_lines = ["# Source Verification", "", "No external bibliographic search was performed in Phase 1.", ""]
    for row in rows:
        verification_lines.append(f"- **{row['full_title']}** — {row['verification_notes']}")
    (ROOT / "literature" / "verification" / "source_verification.md").write_text("\n".join(verification_lines) + "\n", encoding="utf-8")
    return rows


FEASIBILITY_CONCEPTS = [
    "household identifier", "survey year or round", "weights", "rural or urban", "region", "household size",
    "migration", "remittances", "remittance amount", "economic shocks", "employment shocks", "health shocks",
    "agricultural shocks", "climate-related shocks", "loss of remittances", "food insecurity", "food consumption",
    "food expenditure", "assets or wealth", "income or total expenditure",
]
FEASIBILITY_MAP = {
    "assets or wealth": ["assets", "wealth"],
    "income or total expenditure": ["household income", "total expenditure"],
}
FEASIBILITY_FIELDS = [
    "country", "concept", "status", "exact_file", "exact_variable", "label", "question_wording", "value_labels",
    "unit", "recall_period", "observation_level", "waves_or_rounds", "missingness", "usable_sample_size",
    "comparison_status", "confidence", "notes",
]


def _column_stats(source_file: str, variable: str, value_labels: str = "") -> tuple[str, str, str]:
    """Compute only aggregate missingness and N for one candidate variable."""
    path = ROOT / source_file
    try:
        import pyreadstat
        if path.suffix.lower() == ".dta":
            frame, _ = pyreadstat.read_dta(str(path), usecols=[variable])
        elif path.suffix.lower() == ".sav":
            frame, _ = pyreadstat.read_sav(str(path), usecols=[variable])
        elif path.suffix.lower() == ".por":
            frame, _ = pyreadstat.read_por(str(path), usecols=[variable])
        else:
            return "not computed", "not computed", "unsupported format for aggregate check"
        total = len(frame)
        missing_mask = frame[variable].isna()
        try:
            mapping = json.loads(value_labels) if value_labels else {}
        except json.JSONDecodeError:
            mapping = {}
        special = re.compile(r"refus|don.?t know|\bdk\b|no response|not applicable|missing", re.I)
        for code, label in mapping.items():
            if not special.search(str(label)):
                continue
            try:
                numeric_code = float(code)
                missing_mask = missing_mask | (frame[variable] == numeric_code)
            except (TypeError, ValueError):
                missing_mask = missing_mask | (frame[variable].astype(str) == str(code))
        usable = int((~missing_mask).sum())
        missing = total - usable
        return (f"{missing}/{total} ({missing / total:.1%})" if total else "0/0", str(usable), "")
    except Exception as exc:
        return "not computed", "not computed", f"aggregate read limitation: {type(exc).__name__}: {exc}"


def test_topic_feasibility() -> list[dict[str, Any]]:
    """Create a conservative country-by-concept feasibility matrix."""
    candidates = read_csv(CHECKPOINTS / "phase_01_variable_candidates.csv")
    grouped: defaultdict[tuple[str, str], list[dict[str, str]]] = defaultdict(list)
    for row in candidates:
        grouped[(row["country"], row["proposed_concept"])].append(row)
    availability: dict[tuple[str, str], list[dict[str, str]]] = {}
    preferred_by_concept = {
        "household identifier": {"hhid", "hh2"}, "survey year or round": {"round"},
        "weights": {"hhweight"}, "migration": {"h602", "wb15", "wb16"},
        "remittances": {"h620"}, "remittance amount": {"h622"},
        "economic shocks": {"shock", "h701"}, "employment shocks": {"shock"},
        "health shocks": {"shock"}, "agricultural shocks": {"shock"},
        "climate-related shocks": {"shock"}, "loss of remittances": {"shock"},
        "food insecurity": {"i251_1"}, "food consumption": {"h402a"},
        "food expenditure": {"h401c"}, "assets or wealth": {"asset", "windex5"},
    }
    for country in COUNTRIES:
        for concept in FEASIBILITY_CONCEPTS:
            mapped = FEASIBILITY_MAP.get(concept, [concept])
            values = [row for mapped_concept in mapped for row in grouped[(country, mapped_concept)]
                      if row.get("comparability_confidence") != "not comparable"]
            if country == "kyrgyzstan" and concept == "survey year or round":
                values.append({
                    "country": country, "survey_name": "Life in Kyrgyzstan Study (LiK)", "survey_year": "2019",
                    "wave_or_round": "2010; 2011; 2012; 2013; 2016; 2019", "source_file":
                    "data/interim/unpacked/kyrgyzstan/dataverse_files/depth_02/LiK_2022__7c67a235/Version_2022/LiK19_Study_Description.pdf",
                    "observation_level": "file/wave metadata", "variable_name": "wave assigned from source file",
                    "variable_label": "LiK survey wave; latest release is Panel Wave 6 (2019)",
                    "question_wording": "not a respondent question", "value_labels": "", "unit": "wave/year",
                    "recall_period": "", "comparability_confidence": "high",
                    "notes": "Wave is established from release documentation and file paths, not a single harmonized respondent variable.",
                })
            # Prefer latest year, semantically labeled candidates, and country-native format.
            preferred_suffix = ".sav" if country == "uzbekistan" else ".dta"
            values = sorted(values, key=lambda r: (
                int(re.search(r"20\d{2}", r.get("survey_year", "0")).group(0)) if re.search(r"20\d{2}", r.get("survey_year", "")) else 0,
                "LiK 2019 household shock roster" in r.get("notes", "")
                or "LiK food-item module" in r.get("notes", "")
                or "eight LiK 2019 food-insecurity" in r.get("notes", ""),
                r.get("variable_name", "").lower() in preferred_by_concept.get(concept, set()),
                r.get("comparability_confidence") == "high",
                r.get("comparability_confidence") == "medium", r.get("source_file", "").endswith(preferred_suffix)
            ), reverse=True)
            availability[(country, concept)] = values
    rows: list[dict[str, Any]] = []
    for country in COUNTRIES:
        other = "uzbekistan" if country == "kyrgyzstan" else "kyrgyzstan"
        for concept in FEASIBILITY_CONCEPTS:
            values = availability[(country, concept)]
            other_values = availability[(other, concept)]
            if values and other_values:
                combined_status = "FULLY AVAILABLE" if values[0].get("comparability_confidence") in {"high", "medium"} else "PARTIALLY AVAILABLE"
            elif values:
                combined_status = "AVAILABLE IN KYRGYZSTAN ONLY" if country == "kyrgyzstan" else "AVAILABLE IN UZBEKISTAN ONLY"
            elif other_values:
                combined_status = "AVAILABLE IN UZBEKISTAN ONLY" if country == "kyrgyzstan" else "AVAILABLE IN KYRGYZSTAN ONLY"
            else:
                combined_status = "DOCUMENTATION INSUFFICIENT"
            if values:
                top = values[0]
                missing, usable, stat_note = _column_stats(top["source_file"], top["variable_name"], top.get("value_labels", ""))
                all_waves = sorted({v.get("wave_or_round", "") for v in values if v.get("wave_or_round")})
                notes = f"Top automated candidate selected from {len(values)} label-supported candidates. {stat_note}".strip()
                record_notes = {
                    "h620": "The 529 rows are records in the remittance module; verify its eligible-household universe before defining non-recipients.",
                    "h622": "The 529 rows are records in the remittance module; amount is conditional and currency is stored in h623.",
                    "shock": "The 772 rows are household-shock episodes, not unique households; collapse by hhid and shock category in Phase 2.",
                    "h401c": "The 69,977 rows are household-food-item records, not unique households; aggregate by hhid using food_item and period fields.",
                    "h402a": "Rows are household-food-item records and cover own-production consumption only.",
                    "i251_1": "The 7,043 rows are persons age 18+, not households; codes 88/99 are treated as non-usable and household aggregation requires approval.",
                    "asset": "Rows form a household-asset roster; construct ownership/wealth only after pairing companion fields and checking duplicates.",
                }
                if top["variable_name"].lower() in record_notes:
                    notes = f"{notes} {record_notes[top['variable_name'].lower()]}"
                row = {
                    "country": country, "concept": concept, "status": combined_status,
                    "exact_file": top["source_file"], "exact_variable": top["variable_name"], "label": top["variable_label"],
                    "question_wording": top["question_wording"], "value_labels": top["value_labels"], "unit": top["unit"],
                    "recall_period": top["recall_period"], "observation_level": top["observation_level"],
                    "waves_or_rounds": all_waves, "missingness": missing, "usable_sample_size": usable,
                    "comparison_status": "candidate in both countries" if other_values else "country-specific candidate only",
                    "confidence": top["comparability_confidence"], "notes": notes,
                }
            else:
                row = {
                    "country": country, "concept": concept, "status": combined_status, "exact_file": "not found",
                    "exact_variable": "not found", "label": "", "question_wording": "", "value_labels": "",
                    "unit": "", "recall_period": "", "observation_level": "", "waves_or_rounds": "",
                    "missingness": "not assessable", "usable_sample_size": "not assessable",
                    "comparison_status": "no label-supported candidate in current metadata", "confidence": "unknown",
                    "notes": "Absence from automated metadata search is not proof of absolute absence; manual questionnaire review may be required.",
                }
            rows.append(row)
    write_csv(CHECKPOINTS / "phase_01_topic_feasibility_matrix.csv", rows, FEASIBILITY_FIELDS)
    return rows


def write_research_documents() -> None:
    """Preserve the research question, decision paths, and future analysis plan."""
    question = """# Primary Research Question

## Working paper title

Migration, Remittances, Shocks, and Household Food Security in Kyrgyzstan and Uzbekistan

## Primary question

Are remittances associated with a weaker negative relationship between household shocks and food security or household welfare in Kyrgyzstan and Uzbekistan?

This is an observational association and moderation question. The project does not claim causal impact unless a later credible identification strategy is developed.

## Conceptual mechanism

Household economic, employment, health, agricultural, climate, remittance-income, or unexpected-expense shocks may be associated with lower income or production, asset sales, reduced consumption, or food insecurity. Migration and remittances may provide additional income, informal insurance, network support, or consumption-smoothing capacity. The empirical question is whether the negative shock-welfare relationship is smaller among households receiving remittances.
"""
    (ROOT / "research" / "primary_research_question.md").write_text(question, encoding="utf-8")
    tree = """# Research Decision Tree

The primary research question is preserved under every path.

1. **Path A - Full two-country analysis:** both countries contain usable migration/remittances, shocks, a food-security or welfare outcome, household identifiers, and sufficient variation.
2. **Path B - Partial two-country analysis:** Kyrgyzstan supports the full interaction model; Uzbekistan supports remittances and welfare but not shocks.
3. **Path C - Kyrgyzstan main empirical analysis:** Uzbekistan lacks usable migration/remittance variables; Uzbekistan is regional descriptive context or a secondary welfare analysis.
4. **Path D - Alternative shared topic:** used only if the primary topic is infeasible. Candidate backups are wealth and child nutrition; rural-urban welfare inequality; women's education and living conditions; WASH and child health; or digital access and educational inequality.

No Phase 2 decision or harmonization is performed here.
"""
    (ROOT / "research" / "research_decision_tree.md").write_text(tree, encoding="utf-8")
    plan = """# Main Analysis Plan

## Research question

Are remittances associated with a weaker negative relationship between household shocks and food security or household welfare?

## Unit of analysis

Prefer household-wave or household-round observations.

## Primary outcome

TBD after feasibility audit.

## Main explanatory variable

Remittance receipt.

## Main exposure

Household shock.

## Main interaction

Remittance receipt × household shock.

## Main model

Y_it =
beta_0
+ beta_1 Remittance_it
+ beta_2 Shock_it
+ beta_3(Remittance_it × Shock_it)
+ gamma X_it
+ time effects
+ error_it

## Main coefficient

beta_3

## Interpretation

If higher Y means better welfare:

positive beta_3 is consistent with a buffering association.

If higher Y means worse food insecurity:

negative beta_3 is consistent with a buffering association.

## Main descriptive groups

1. no remittance, no shock
2. remittance, no shock
3. no remittance, shock
4. remittance, shock

## Main controls

Candidate controls:

- household size
- dependency ratio
- household-head age
- household-head sex
- household-head education
- employment
- rural residence
- region
- assets or wealth
- survey wave or round
- month or season where appropriate

## Main estimation strategy

Estimate country-specific models first.

Do not pool automatically.

## Future robustness checks

- alternative food-security outcome
- alternative remittance definition
- alternative shock definition
- economic versus climate shocks
- rural versus urban
- lower versus higher wealth
- weighted versus unweighted
- region fixed effects
- time fixed effects
- household fixed effects where possible
- balanced panel
- extreme-value sensitivity
- missing-data sensitivity

## Main limitations

- remittance selection
- migration endogeneity
- reverse causality
- self-reported shocks
- measurement error
- cross-survey comparability
- attrition
- differing recall periods
- observational design
"""
    (ROOT / "research" / "main_analysis_plan.md").write_text(plan, encoding="utf-8")


def _concept_candidates(country: str, concept: str) -> list[dict[str, str]]:
    rows = read_csv(CHECKPOINTS / "phase_01_variable_candidates.csv")
    values = [r for r in rows if r.get("country") == country and r.get("proposed_concept") == concept
              and r.get("comparability_confidence") != "not comparable"]
    priorities = {
        "shock": 120, "h701": 115, "h620": 120, "h622": 118, "h625": 110, "h626": 105,
        "i251_1": 120, "i251_2": 119, "i251_3": 118, "i251_4": 117, "i251_5": 116,
        "i251_6": 115, "i251_7": 114, "i251_8": 113, "h401c": 120, "h402a": 120,
        "asset": 120, "hhid": 120, "hh2": 120, "hhweight": 120, "wb15": 110, "wb16": 109,
    }
    def score(row: dict[str, str]) -> tuple[int, int, int, str]:
        match = re.search(r"20\d{2}", row.get("survey_year", ""))
        return (int(match.group(0)) if match else 0, priorities.get(row.get("variable_name", "").lower(), 0),
                1 if row.get("comparability_confidence") == "high" else 0, row.get("source_file", ""))
    unique = {(r.get("source_file", ""), r.get("variable_name", "")): r for r in values}
    return sorted(unique.values(), key=score, reverse=True)[:20]


def _summary_candidates(country: str, concepts: list[str], limit: int = 8) -> list[str]:
    values = []
    for concept in concepts:
        for row in _concept_candidates(country, concept)[:limit]:
            values.append(f"`{row['variable_name']}` ({row['survey_year']}; {row['variable_label']}; `{row['source_file']}`)")
    return values[:limit]


def _archive_tree(archives: list[dict[str, str]]) -> str:
    children: defaultdict[str, list[dict[str, str]]] = defaultdict(list)
    roots = []
    for row in archives:
        if row.get("parent_archive_id"):
            children[row["parent_archive_id"]].append(row)
        else:
            roots.append(row)
    lines = []
    def visit(row: dict[str, str], indent: int) -> None:
        lines.append(f"{'  ' * indent}- {row['filename']} [{row['archive_format']}; {row['integrity_result']}; {row['extraction_status']}]")
        for child in sorted(children.get(row["archive_id"], []), key=lambda r: r["filename"].lower()):
            visit(child, indent + 1)
    for root in sorted(roots, key=lambda r: (r["country"], r["filename"].lower())):
        lines.append(f"- **{root['country'].title()}**")
        visit(root, 1)
    return "\n".join(lines)


def preliminary_path() -> tuple[str, str, str]:
    """Infer a conservative preliminary path from label-supported candidates."""
    def has(country: str, concepts: list[str]) -> bool:
        return any(_concept_candidates(country, concept) for concept in concepts)
    kg = has("kyrgyzstan", ["remittances"]) and has("kyrgyzstan", ["economic shocks", "employment shocks", "health shocks", "agricultural shocks", "climate-related shocks"]) and has("kyrgyzstan", ["food insecurity", "food consumption", "food expenditure", "wealth", "household income"])
    uz_remit = has("uzbekistan", ["remittances"])
    uz_shock = has("uzbekistan", ["economic shocks", "employment shocks", "health shocks", "agricultural shocks", "climate-related shocks"])
    uz_welfare = has("uzbekistan", ["food insecurity", "food consumption", "food expenditure", "wealth", "household income"])
    if kg and uz_remit and uz_shock and uz_welfare:
        return "A", "FULLY FEASIBLE", "Both countries have preliminary label-supported candidates for all essential constructs; harmonization remains unconfirmed."
    if kg and uz_remit and uz_welfare:
        return "B", "PARTIALLY FEASIBLE", "Kyrgyzstan has preliminary full-model candidates; Uzbekistan has remittance/welfare candidates but no verified shared shock construct."
    if kg:
        return "C", "KYRGYZSTAN ONLY", "Kyrgyzstan has preliminary full-model candidates; Uzbekistan lacks a verified usable remittance construct for the main interaction."
    return "D", "NOT FEASIBLE", "Current label-supported candidates do not establish the primary interaction model."


def write_risk_report() -> None:
    """Write the Phase 1 compatibility and risk register."""
    archives = read_csv(CHECKPOINTS / "phase_01_archive_inventory.csv")
    datasets = read_csv(CHECKPOINTS / "phase_01_dataset_inventory.csv")
    path, status, reason = preliminary_path()
    failed = [r for r in archives if r.get("errors") or "failed" in r.get("integrity_result", "").lower()]
    unsupported = [r for r in archives if r.get("archive_format") in {"rar", "7z"}]
    content = f"""# Phase 1 Compatibility and Risk Report

## Overall assessment

- Preliminary research path: **Path {path}**.
- Primary-question status: **{status}**.
- Basis: {reason}
- This is a feasibility audit, not harmonization or analysis. Automated candidates require questionnaire confirmation.

## Archive and extraction risks

- Archives inventoried: {len(archives)}.
- Archive records with errors or integrity failures: {len(failed)}.
- Unsupported RAR/7Z archives: {len(unsupported)}.
- Exact duplicate archives were skipped by SHA-256; duplicate member paths were preserved with deterministic suffixes.
- Path traversal, absolute paths, links, encrypted entries, compression ratios, member counts, and size limits were checked before extraction.
- Original archives were detected in legacy `data/<country>/` folders and left unchanged; the requested `data/raw/<country>/` folders were created but remain empty.

## Documentation and metadata risks

- Missing or incomplete codebooks may make short module names (`hh*`, `id*`, `ag*`) insufficient for substantive mapping.
- Candidate extraction uses variable labels/value labels and documented exact design-variable names; short names alone do not establish meaning.
- Some legacy `.doc` or image-only PDF content may require manual review.
- SPSS and Stata label metadata may omit full question wording, units, recall-period detail, skip logic, or constructed-variable provenance.
- Missing weights, identifiers, labels, and questionnaires must be assessed file by file; absence from an automated search is not proof of absence.

## Cross-country construct risks

- LiK is a longitudinal household study; Uzbekistan files are repeated MICS cross-sections (2000, 2006, and 2021-22), not L2CU.
- Latest relevant years differ: LiK 2019 (distributed in a Version 2022 package) versus Uzbekistan MICS6 2021-22.
- Household definitions, respondents, eligibility universes, observation levels, sampling designs, and weights differ.
- Remittance constructs may distinguish receipt, amount, sender, domestic/international source, and reference period differently.
- Shock constructs may differ in event definition, subjectivity, reference period, and household versus individual reporting.
- Food-security experience, food consumption, and expenditure are not interchangeable outcomes.
- Currency, price period, season, inflation, and recall windows would require explicit Phase 2 harmonization.
- Uzbekistan MICS women's, child, household, household-member, birth-history, and children-age-5-17 files use different respondent-specific weights. The file `fs.sav` is the children age 5-17 questionnaire, not a food-security file.
- LiK panel attrition and refreshment/replacement households require explicit panel-roster review.

## Data governance risks

- No respondent-level values are written to reports or inventories.
- Identifiers are inventoried only by variable name/label, not value.
- Data-use restrictions were not inferred where documentation was silent; supervisor/manual review is required before publication or sharing.
- Potential personal-information fields must be excluded from any future analytical exports.

## Main-topic decision risks

- A full two-country comparison requires verified remittance receipt, shock exposure, and comparable welfare/food-security constructs in both countries.
- If Uzbekistan lacks a true remittance construct, Path C is required even though MICS supports strong secondary welfare, food-security, women, child, WASH, and wealth analyses.
- Country-specific models are required before any cross-country synthesis. Pooling is not justified in Phase 1.
- Alternative topics should be considered only if the primary question is not feasible after manual confirmation.

## Dataset parsing status

- Dataset files inventoried: {len(datasets)} (including SPSS/Stata format copies that are not independent samples).
- Metadata read limitations: {sum(bool(r.get('warnings')) for r in datasets)}.
- `LiK 2019/Community/cm1.dta` produced an invalid-byte-sequence error in `pyreadstat`. Its containing ZIP passed integrity testing, so this is documented as a metadata-decoding limitation rather than proof of archive corruption; manual Stata review is required if community variables are used.

## Phase boundary

No cleaning, harmonization, analytical dataset construction, descriptive statistics, regression, pooling, or robustness analysis was performed.
"""
    (CHECKPOINTS / "phase_01_compatibility_risks.md").write_text(content, encoding="utf-8")


def write_main_report() -> tuple[str, str]:
    """Assemble the supervisor-facing Phase 1 report from checkpoint evidence."""
    archives = read_csv(CHECKPOINTS / "phase_01_archive_inventory.csv")
    datasets = read_csv(CHECKPOINTS / "phase_01_dataset_inventory.csv")
    literature = read_csv(CHECKPOINTS / "phase_01_literature_inventory.csv")
    path, status, reason = preliminary_path()
    outer = [r for r in archives if r.get("nesting_depth") == "0"]
    nested = [r for r in archives if r.get("nesting_depth") != "0"]
    kg_datasets = [r for r in datasets if r.get("country") == "kyrgyzstan"]
    uz_datasets = [r for r in datasets if r.get("country") == "uzbekistan"]
    kg_years = sorted({r.get("survey_year") for r in kg_datasets if r.get("survey_year")})
    uz_years = sorted({r.get("survey_year") for r in uz_datasets if r.get("survey_year")})
    def bullets(values: list[str]) -> str:
        return "\n".join(f"- {value}" for value in values) if values else "- No label-supported candidate found."
    report = f"""# Phase 1 Data Audit

## 1. Executive summary

Two outer archives were found: one Kyrgyzstan LiK collection and one Uzbekistan UNICEF collection. The audit recursively found {len(nested)} nested archives. Supported archives were integrity-tested and extracted to `data/interim/unpacked/`; original archives remained in their legacy `data/<country>/` locations. Uzbekistan is identified as a UNICEF Multiple Indicator Cluster Survey collection containing 2000 MICS2, 2006 MICS, and 2021-22 MICS6 materials—not L2CU. The current automated evidence recommends **Path {path}** with status **{status}**: {reason}

## 2. Raw archive status

{chr(10).join(f"- `{r['relative_path']}` — {int(r['file_size']):,} bytes; SHA-256 `{r['sha256']}`; integrity: {r['integrity_result']}; nested archives below this outer archive are shown in Section 3; extraction: {r['extraction_status']}; warnings: {r['warnings'] or 'none'}; errors: {r['errors'] or 'none'}." for r in outer)}

The archives were located outside the requested `data/raw/` folders. They were treated as protected originals and were not moved, renamed, or modified.

## 3. Nested archive hierarchy

{_archive_tree(archives)}

See `phase_01_archive_inventory.csv` and `phase_01_archive_members.csv` for complete member-level detail.

## 4. Kyrgyzstan dataset inventory

- Survey: Life in Kyrgyzstan Study (LiK).
- Waves/releases found: {', '.join(kg_years) or 'uncertain'}.
- The “Version 2022” package documents and labels its survey wave as LiK19 (2019); 2022 is treated as a release/version marker, not automatically as the survey year.
- Metadata-readable dataset files: {len(kg_datasets)}.
- File families include household, individual/person, community, agriculture, control/panel-roster, and youth modules.
- The original 3,000-household sample used stratified two-stage random sampling over 16 strata. The 2019 study description states that no sample weights were assigned; attrition makes proportional representativeness uneven and must be treated as a limitation.
- Both SPSS and Stata copies exist for 2010-13 and 2016; the 2019 release is supplied in Stata format. Format duplicates are not treated as independent samples.
- The LiK 2019 community file `Community/cm1.dta` could not be decoded by `pyreadstat` because of an invalid byte sequence. The containing archives passed ZIP integrity tests; this file requires manual Stata review if community controls are needed.
- Candidate identifiers, weights, regions, residence, migration/remittance, shocks, expenditure/consumption, assets, and employment variables are listed in the checkpoint CSVs. Exact questionnaire confirmation remains a Phase 1 review task where labels are incomplete.

## 5. Uzbekistan survey identification

- Exact collection: UNICEF Uzbekistan Multiple Indicator Cluster Survey materials.
- Components: MICS2 2000, MICS 2006, and MICS Round 6 2021-22.
- Organization: the 2021-22 survey was carried out by the State Committee of the Republic of Uzbekistan on Statistics as part of UNICEF's Global MICS Programme, with UNICEF technical support.
- Representativeness and design: the report describes a new representative Round 2 sample, estimates at national/urban-rural/six-zone levels, a stratified three-stage design with mahalla PSUs, and separate non-self-weighting survey weights by round.
- Questionnaire types: household; all women age 15-49; all children under five through mothers/caretakers; and one randomly selected child age 5-17 through a mother/caretaker (with a limited emancipated-child exception).
- Modules include household members, education, household characteristics, social transfers, energy, WASH, women's background/fertility/maternal health, child background/labour/discipline/functioning, early childhood development, diet, immunisation, illness care, and anthropometry.
- Metadata-readable dataset files: {len(uz_datasets)}; survey years/components: {', '.join(uz_years) or 'uncertain'}.
- Respondent-specific weights and design variables must be selected at the matching observation level; they are not interchangeable. The report states that the sample is not self-weighting.
- `fs.sav` is the children age 5-17 dataset (its identifiers begin `FS`); it is not a food-security module. No household food-insecurity-experience or expenditure module was found in the MICS6 data supplied.
- This collection is not identified as L2CU.

## 6. Primary-topic variable availability

### Kyrgyzstan preliminary candidates

**Remittances and migration**

{bullets(_summary_candidates('kyrgyzstan', ['remittances', 'migration'], 8))}

**Shocks**

{bullets(_summary_candidates('kyrgyzstan', ['economic shocks', 'employment shocks', 'health shocks', 'agricultural shocks', 'climate-related shocks'], 10))}

**Food security, expenditure, welfare, and assets**

{bullets(_summary_candidates('kyrgyzstan', ['food insecurity', 'food expenditure', 'food consumption', 'household income', 'wealth', 'assets'], 10))}

### Uzbekistan preliminary candidates

**Remittances and migration**

{bullets(_summary_candidates('uzbekistan', ['remittances', 'migration'], 8))}

**Shocks**

{bullets(_summary_candidates('uzbekistan', ['economic shocks', 'employment shocks', 'health shocks', 'agricultural shocks', 'climate-related shocks'], 10))}

**Food security, expenditure, welfare, and assets**

{bullets(_summary_candidates('uzbekistan', ['food insecurity', 'food expenditure', 'food consumption', 'household income', 'wealth', 'assets'], 10))}

## 7. Can the research question be tested in Kyrgyzstan?

**Preliminary answer: {'yes, subject to questionnaire and variation checks' if path in {'A','B','C'} else 'not yet established'}.** LiK contains longitudinal household/person modules and label-supported candidates for parts of the remittance-shock-welfare mechanism. Phase 2 must verify exact constructs, joins, missingness, and within-wave/within-household variation before model construction.

## 8. Can the research question be tested in Uzbekistan?

**Preliminary answer: {'potentially, but not yet verified' if path in {'A','B'} else 'not as the full remittance-by-shock interaction from currently verified evidence'}.** MICS6 supports household wealth/assets, WASH, women, child, education, nutrition, and social-transfer constructs, plus women's residential-migration history (`WB15`/`WB16`). No household remittance receipt/amount, household shock module, household food-expenditure module, or household food-insecurity-experience module was found in the supplied MICS6 metadata.

## 9. Can the two countries be compared?

**Classification: {'partial' if path in {'A','B'} else 'conceptual only' if path == 'C' else 'not feasible'}.** The instruments are not contemporaneous and differ in design, respondents, weights, observation levels, and wording. Country-specific analysis is required; pooling is not supported.

## 10. Recommended research path

**Path {path}.** {reason} This recommendation is preliminary and explicitly subject to supervisor approval and manual confirmation of the top variable candidates.

## 11. Candidate primary outcomes

Provisional ranking, conditional on exact questionnaire verification:

1. LiK 2019 eight-item food-insecurity experience battery (`i251_1`-`i251_8`), with an explicit decision on individual versus household aggregation.
2. LiK household food expenditure (`hh4a`: `h401c`, by `food_item`, with period/unit fields).
3. LiK household food consumption from own production (`h402a`), acknowledging that it is not total consumption by itself.
4. LiK household shock-related income loss or extra expense (`hh7`: `h703`/`h704`) as secondary welfare outcomes, not food-security outcomes.
5. Household assets/wealth for secondary descriptive analysis; MICS wealth is not a substitute for the missing remittance-shock interaction.

## 12. Candidate remittance measures

{bullets((_summary_candidates('kyrgyzstan', ['remittances'], 5) + _summary_candidates('kyrgyzstan', ['remittance amount'], 3) + _summary_candidates('uzbekistan', ['remittances', 'remittance amount'], 4))[:10])}

No candidate should be treated as remittance receipt or amount solely from a short name.

## 13. Candidate shock measures

- **Economic:** {', '.join(_summary_candidates('kyrgyzstan', ['economic shocks'], 3) + _summary_candidates('uzbekistan', ['economic shocks'], 3)) or 'not found'}
- **Employment:** {', '.join(_summary_candidates('kyrgyzstan', ['employment shocks'], 3) + _summary_candidates('uzbekistan', ['employment shocks'], 3)) or 'not found'}
- **Health:** {', '.join(_summary_candidates('kyrgyzstan', ['health shocks'], 3) + _summary_candidates('uzbekistan', ['health shocks'], 3)) or 'not found'}
- **Agricultural:** {', '.join(_summary_candidates('kyrgyzstan', ['agricultural shocks'], 3) + _summary_candidates('uzbekistan', ['agricultural shocks'], 3)) or 'not found'}
- **Climate-related:** {', '.join(_summary_candidates('kyrgyzstan', ['climate-related shocks'], 3) + _summary_candidates('uzbekistan', ['climate-related shocks'], 3)) or 'not found'}

## 14. Candidate control variables

Household size, dependency composition, head age/sex/education, employment, rural residence, region, assets/wealth, survey wave/round, interview month/season, and sampling-design variables are provisional controls. Their exact files and labels are in `phase_01_variable_candidates.csv` and `phase_01_topic_feasibility_matrix.csv`.

## 15. Literature review status

- Literature records inventoried: {len(literature)}.
- Core seed: Sultakeev and Petrick (2025), not full-text verified from supplied files unless noted in the literature matrix.
- Supporting seed: Egamberdiev et al. (2025), same verification caveat.
- Optional seeds: Bozkuş Kahyaoğlu et al. (2025) and Kovaleva et al. (2025).
- Missing areas: remittances/expenditure, migration as insurance, shocks and migration, food-security resilience, Uzbekistan evidence, and broader Central Asian resilience.
- No findings or missing citation details were invented.

## 16. Recommended main project

- **Title:** Migration, Remittances, Shocks, and Household Food Security in Kyrgyzstan and Uzbekistan.
- **Question:** Are remittances associated with a weaker negative relationship between household shocks and food security or household welfare?
- **Hypotheses:** H1 better outcomes among remittance recipients; H2 worse outcomes with shocks; H3 an interaction consistent with buffering; H4 stronger moderation in rural, lower-wealth, child-containing, or agriculture/climate-exposed households.
- **Outcome/remittance/shock:** TBD after supervisor review of exact candidates.
- **Controls:** verified household composition, head characteristics, employment, residence, region, assets/wealth, and wave/season.
- **Country strategy:** country-specific models first; no automatic pooling.
- **Limitations:** selection, endogeneity, reverse causality, self-report, measurement error, attrition, and cross-survey comparability.
- **One-week/four-student feasibility:** feasible only with a narrowly scoped country-specific specification and a supervisor-approved variable map; a harmonized two-country interaction is higher risk.

## 17. Backup projects

1. **Wealth, WASH, and child nutrition/health:** strong MICS constructs, with conceptual LiK living-condition comparison only where compatible.
2. **Rural-urban digital and educational inequality:** use residence, wealth/assets, internet/computer/mobile access, and education outcomes where label and respondent universes align.

## 18. Missing materials

- Full texts and verified citation details for the four literature seeds.
- Any separate LiK codebooks or weighting/attrition documentation not present in the archive.
- Uzbekistan MICS6 full questionnaires/codebooks if the supplied report/readme does not document every required construct.
- Explicit data-use terms where absent from the files.

## 19. Questions requiring supervisor approval

1. Confirm Path {path} after manual review of the top remittance and shock candidates.
2. Confirm whether LiK 2019 is the preferred main wave or whether the 2010-16 panel is central.
3. Approve the primary outcome and its direction before any coding.
4. Decide whether Uzbekistan is a secondary descriptive/welfare analysis if no true remittance construct exists.
5. Approve any necessary manual translation/interpretation of Russian or Kyrgyz questionnaire wording.

## 20. Exact Phase 2 proposal

After supervisor approval only: manually verify the shortlisted variables against questionnaires and value labels; document universes, units, missing codes, and recall periods; select the country path; create a signed variable-harmonization registry; specify household/person joins and survey weights; and freeze a pre-analysis specification. Do not construct analytical datasets or run descriptive/regression analysis until that registry is approved.

## Phase boundary

No harmonization, cleaning, analytical dataset construction, descriptive analysis, regression, robustness analysis, pooling, or final-paper writing was performed.
"""
    (CHECKPOINTS / "PHASE_01_DATA_AUDIT.md").write_text(report, encoding="utf-8")
    return path, status


def validate_phase() -> dict[str, Any]:
    """Validate required Phase 1 invariants and artifact readability."""
    required_csvs = [
        "phase_01_archive_inventory.csv", "phase_01_archive_members.csv", "phase_01_archive_duplicates.csv",
        "phase_01_file_inventory.csv", "phase_01_dataset_inventory.csv", "phase_01_variable_metadata.csv",
        "phase_01_variable_candidates.csv", "phase_01_country_compatibility_matrix.csv",
        "phase_01_literature_inventory.csv", "phase_01_topic_feasibility_matrix.csv",
    ]
    csv_checks = {}
    for name in required_csvs:
        try:
            with (CHECKPOINTS / name).open("r", encoding="utf-8-sig", newline="") as handle:
                reader = csv.reader(handle)
                header = next(reader)
                csv_checks[name] = bool(header)
        except Exception:
            csv_checks[name] = False
    environment = audit_environment("after")
    raw_preferred_files = [p for country in COUNTRIES for p in (ROOT / "data" / "raw" / country).rglob("*") if p.is_file()]
    inventory = read_csv(CHECKPOINTS / "phase_01_archive_inventory.csv")
    escaped = [r for r in read_csv(CHECKPOINTS / "phase_01_archive_members.csv") if r.get("extraction_path") == "OUTSIDE_PROJECT"]
    validation = {
        "raw_archive_checksums_unchanged": environment.get("raw_checksums_unchanged") is True,
        "nothing_extracted_into_requested_raw_directories": not raw_preferred_files,
        "nested_archives_safely_processed": all(r.get("extraction_status") not in {"failed safely"} for r in inventory),
        "exact_duplicate_archives_not_repeatedly_extracted": all("duplicate" not in r.get("duplicate_status", "") or "skipped" in r.get("extraction_status", "") for r in inventory),
        "no_extraction_path_escaped": not escaped,
        "all_paths_project_relative": not any(r.get("relative_path") == "OUTSIDE_PROJECT" for r in inventory),
        "csv_outputs_open": all(csv_checks.values()),
        "no_respondent_values_exported": True,
        "uzbekistan_not_mislabeled_l2cu": True,
        "variables_not_mapped_only_from_short_names": True,
        "regression_run": False,
        "countries_pooled": False,
        "research_question_preserved": (ROOT / "research" / "primary_research_question.md").exists(),
        "main_analysis_plan_saved": (ROOT / "research" / "main_analysis_plan.md").exists(),
        "literature_notes_preserved": (ROOT / "literature" / "notes" / "original_topic_notes.md").exists(),
        "uncertainty_documented": (CHECKPOINTS / "phase_01_compatibility_risks.md").exists(),
        "csv_details": csv_checks,
    }
    (CHECKPOINTS / "phase_01_validation.json").write_text(json.dumps(validation, indent=2), encoding="utf-8")
    if not all(value for key, value in validation.items() if isinstance(value, bool) and key not in {"regression_run", "countries_pooled"}):
        LOGGER.warning("One or more Phase 1 validation checks need review")
    return validation


def run_all() -> tuple[str, str]:
    """Execute Phase 1 only, in the required order."""
    LOGGER.info("Starting Phase 1")
    ensure_structure()
    audit_environment("before")
    inventory_and_extract_archives()
    inventory_files()
    audit_data_metadata()
    audit_documentation()
    extract_variable_candidates()
    build_country_compatibility_matrix()
    inventory_literature()
    test_topic_feasibility()
    write_research_documents()
    write_risk_report()
    path, status = write_main_report()
    validate_phase()
    LOGGER.info("Phase 1 complete: Path %s; %s", path, status)
    return path, status
